# api/document_reader.py
"""
Unified document extractor supporting PDF, DOCX, and TXT files for Job Descriptions and Resumes.
"""

from __future__ import annotations
import os
from pathlib import Path
from typing import Optional


def extract_text_from_file(file_path: str | Path) -> str:
    """
    Extract raw text from a document based on its extension (.pdf, .docx, .doc, .txt, .md).
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(path)
    elif ext in (".txt", ".md", ".rtf"):
        return _extract_from_text(path)
    else:
        # Fallback to text reading
        return _extract_from_text(path)


def _extract_from_pdf(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        return "\n".join(pages_text).strip()
    except Exception:
        # Fallback to existing parser
        try:
            from parser.textAndLinkSeperator import extract_pdf
            res = extract_pdf(str(path))
            return res.get("text", "").strip()
        except Exception as e:
            raise RuntimeError(f"Failed to read PDF file: {e}")


def _extract_from_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(path))
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines).strip()
    except Exception as e:
        # Fallback to binary/plain text scan
        try:
            return _extract_from_text(path)
        except Exception:
            raise RuntimeError(f"Failed to read Word document (.docx): {e}")


def _extract_from_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc, errors="replace") as f:
                return f.read().strip()
        except Exception:
            continue
    raise RuntimeError(f"Could not read text file: {path.name}")
